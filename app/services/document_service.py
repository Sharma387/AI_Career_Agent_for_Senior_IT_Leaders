import io
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from jinja2 import Environment, FileSystemLoader

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"


def _parse_resume_text(text: str) -> dict:
    """Parse raw resume text into structured sections."""
    sections = {
        "summary": "",
        "experience": [],
        "education": [],
        "skills": {},
        "certifications": [],
    }

    if not text:
        return sections

    lines = [l.strip() for l in text.split("\n") if l.strip()]
    current_section = "summary"
    current_item = None

    for line in lines:
        lower = line.lower()

        if any(kw in lower for kw in ["experience", "employment", "work history"]):
            current_section = "experience"
            current_item = None
            continue
        elif any(kw in lower for kw in ["education", "academic", "qualifications"]):
            current_section = "education"
            current_item = None
            continue
        elif any(kw in lower for kw in ["skills", "competencies", "technical"]):
            current_section = "skills"
            current_item = None
            continue
        elif any(kw in lower for kw in ["certifications", "certificates", "licenses"]):
            current_section = "certifications"
            current_item = None
            continue
        elif any(kw in lower for kw in ["summary", "profile", "about", "objective"]):
            current_section = "summary"
            current_item = None
            continue

        if current_section == "summary":
            sections["summary"] += line + " "
        elif current_section == "experience":
            if not current_item or (len(line) < 80 and not line.startswith("-") and not line.startswith("•")):
                if current_item:
                    sections["experience"].append(current_item)
                current_item = {"title": line, "company": "", "dates": "", "bullets": [], "description": ""}
            else:
                if current_item:
                    clean = line.lstrip("-•· ")
                    current_item["bullets"].append(clean)
        elif current_section == "education":
            if not current_item:
                current_item = {"degree": line, "institution": "", "year": ""}
            else:
                current_item["institution"] = line
                sections["education"].append(current_item)
                current_item = None
        elif current_section == "skills":
            if ":" in line:
                cat, skills = line.split(":", 1)
                sections["skills"][cat.strip()] = [s.strip() for s in skills.split(",") if s.strip()]
            else:
                sections["skills"]["General"] = sections["skills"].get("General", []) + [line]
        elif current_section == "certifications":
            sections["certifications"].append({"name": line, "issuer": ""})

    if current_item and current_section == "experience":
        sections["experience"].append(current_item)

    sections["summary"] = sections["summary"].strip()
    return sections


def _parse_cover_letter_text(text: str) -> dict:
    """Parse cover letter text into structured data."""
    if not text:
        return {"body_paragraphs": []}

    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    body = []
    for p in paragraphs:
        lower = p.lower()
        if lower.startswith("dear") or lower.startswith("yours") or lower.startswith("sincerely"):
            continue
        if "thank you" in lower and len(body) > 0:
            continue
        body.append(p)

    return {"body_paragraphs": body if body else [text]}


def render_resume_html(profile_data: dict, job_data: dict = None) -> str:
    """Render resume as styled HTML using template."""
    env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
    template = env.get_template("resume.html")

    resume = _parse_resume_text(profile_data.get("resume_text", ""))

    context = {
        "candidate_name": profile_data.get("full_name", "Candidate"),
        "headline": profile_data.get("headline", ""),
        "email": profile_data.get("email", ""),
        "phone": profile_data.get("phone", ""),
        "location": profile_data.get("location", ""),
        "linkedin": profile_data.get("linkedin", ""),
        "summary": profile_data.get("summary") or resume["summary"],
        "experience": resume["experience"],
        "projects": profile_data.get("projects", []),
        "skills": profile_data.get("skills") or resume["skills"],
        "education": resume["education"],
        "certifications": profile_data.get("certifications", []) or resume["certifications"],
    }

    if job_data:
        context["headline"] = job_data.get("target_role", context.get("headline", ""))

    return template.render(**context)


def render_cover_letter_html(profile_data: dict, job_data: dict, cover_letter_text: str) -> str:
    """Render cover letter as styled HTML using template."""
    env = Environment(loader=FileSystemLoader(str(TEMPLATES_DIR)))
    template = env.get_template("cover_letter.html")

    parsed = _parse_cover_letter_text(cover_letter_text)

    context = {
        "candidate_name": profile_data.get("full_name", "Candidate"),
        "email": profile_data.get("email", ""),
        "phone": profile_data.get("phone", ""),
        "location": profile_data.get("location", ""),
        "date": datetime.now().strftime("%d %B %Y"),
        "company_name": job_data.get("company", "the company"),
        "recipient_name": job_data.get("recipient_name", ""),
        "recipient_title": job_data.get("recipient_title", ""),
        "company_address": job_data.get("company_address", ""),
        "body_paragraphs": parsed["body_paragraphs"],
    }

    return template.render(**context)


def generate_resume_docx(profile_data: dict, job_data: dict = None) -> bytes:
    """Generate a professional NZ-style resume as DOCX."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    resume = _parse_resume_text(profile_data.get("resume_text", ""))

    name = profile_data.get("full_name", "Candidate")
    h = doc.add_heading(name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        run.font.size = Pt(20)

    contact_parts = []
    if profile_data.get("email"):
        contact_parts.append(profile_data["email"])
    if profile_data.get("phone"):
        contact_parts.append(profile_data["phone"])
    if profile_data.get("location"):
        contact_parts.append(profile_data["location"])
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.style.font.size = Pt(9)
        p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph("")

    if profile_data.get("summary") or resume["summary"]:
        h = doc.add_heading("Professional Summary", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            run.font.size = Pt(11)
        doc.add_paragraph(profile_data.get("summary") or resume["summary"])

    if resume["experience"]:
        h = doc.add_heading("Experience", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            run.font.size = Pt(11)
        for exp in resume["experience"]:
            p = doc.add_paragraph()
            run = p.add_run(exp["title"])
            run.bold = True
            if exp.get("company"):
                p.add_run(f" — {exp['company']}")
            if exp.get("dates"):
                p.add_run(f"  ({exp['dates']})")
            for bullet in exp.get("bullets", []):
                bp = doc.add_paragraph(bullet, style="List Bullet")

    if profile_data.get("projects"):
        h = doc.add_heading("Key Projects", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            run.font.size = Pt(11)
        for proj in profile_data["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("title", ""))
            run.bold = True
            if proj.get("role"):
                p.add_run(f" — {proj['role']}")
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            if proj.get("technologies"):
                doc.add_paragraph(f"Technologies: {proj['technologies']}")

    if profile_data.get("skills") or resume["skills"]:
        h = doc.add_heading("Skills", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            run.font.size = Pt(11)
        skills = profile_data.get("skills") or resume["skills"]
        if isinstance(skills, dict):
            for cat, skill_list in skills.items():
                p = doc.add_paragraph()
                run = p.add_run(f"{cat}: ")
                run.bold = True
                p.add_run(", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list))
        elif isinstance(skills, list):
            doc.add_paragraph(", ".join(str(s) for s in skills))

    if resume["education"]:
        h = doc.add_heading("Education", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            run.font.size = Pt(11)
        for edu in resume["education"]:
            p = doc.add_paragraph()
            run = p.add_run(edu["degree"])
            run.bold = True
            if edu.get("institution"):
                p.add_run(f" — {edu['institution']}")
            if edu.get("year"):
                p.add_run(f" ({edu['year']})")

    if profile_data.get("certifications"):
        h = doc.add_heading("Certifications", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
            run.font.size = Pt(11)
        for cert in profile_data["certifications"]:
            name = cert.get("name", "") if isinstance(cert, dict) else str(cert)
            issuer = cert.get("issuer", "") if isinstance(cert, dict) else ""
            doc.add_paragraph(f"{name} — {issuer}" if issuer else name, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def generate_cover_letter_docx(profile_data: dict, job_data: dict, cover_letter_text: str) -> bytes:
    """Generate a professional cover letter as DOCX."""
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

    name = profile_data.get("full_name", "Candidate")
    h = doc.add_heading(name, level=0)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x52, 0x76)
        run.font.size = Pt(16)

    contact_parts = []
    if profile_data.get("email"):
        contact_parts.append(profile_data["email"])
    if profile_data.get("phone"):
        contact_parts.append(profile_data["phone"])
    if contact_parts:
        p = doc.add_paragraph(" | ".join(contact_parts))
        p.paragraph_format.space_after = Pt(12)

    doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    doc.add_paragraph("")

    if job_data.get("recipient_name"):
        doc.add_paragraph(job_data["recipient_name"])
        if job_data.get("recipient_title"):
            doc.add_paragraph(job_data["recipient_title"])
        doc.add_paragraph(job_data.get("company", ""))
        doc.add_paragraph("")

    parsed = _parse_cover_letter_text(cover_letter_text)

    doc.add_paragraph(f"Dear {job_data.get('recipient_name', 'Hiring Manager')},")

    for para in parsed["body_paragraphs"]:
        doc.add_paragraph(para)

    doc.add_paragraph("")
    doc.add_paragraph("Yours sincerely,")
    doc.add_paragraph(name)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
